"""
Advanced Document Processing Service
Handles AI-powered Q&A, smart editing, format conversion, and structured JSON output
"""

import os
import json
import io
import tempfile
import logging
import re
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
import uuid

# File processing imports
import pdfplumber
from docx import Document as DocxDocument
from docx.shared import Inches
import pytesseract
from PIL import Image
import fitz  # PyMuPDF
import pandas as pd
from openpyxl import Workbook
import csv

# Format conversion imports
import markdown
import html2text
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from fpdf import FPDF

# AI and web imports
import requests
from bs4 import BeautifulSoup

# Database imports
from .extensions import db
from .models import Document, ProcessingJob

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Advanced document processing with AI capabilities"""
    
    def __init__(self, gemini_api_key: str):
        self.gemini_api_key = gemini_api_key
        self.gemini_api_url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-pro:generateContent"
        
    def extract_enhanced_text(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Enhanced text extraction with metadata"""
        try:
            result = {
                'text': '',
                'metadata': {},
                'structure': {},
                'success': True,
                'error': None
            }
            
            if file_type.lower() == 'pdf':
                result = self._extract_pdf_enhanced(file_path)
            elif file_type.lower() in ['docx', 'doc']:
                result = self._extract_docx_enhanced(file_path)
            elif file_type.lower() in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff']:
                result = self._extract_image_enhanced(file_path)
            else:
                result['success'] = False
                result['error'] = f"Unsupported file type: {file_type}"
                
            return result
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return {
                'text': '',
                'metadata': {},
                'structure': {},
                'success': False,
                'error': str(e)
            }
    
    def _extract_pdf_enhanced(self, file_path: str) -> Dict[str, Any]:
        """Enhanced PDF extraction with structure analysis"""
        text = ""
        metadata = {}
        structure = {'pages': [], 'tables': [], 'images': []}
        
        # Using PyMuPDF for enhanced extraction
        doc = fitz.open(file_path)
        metadata = {
            'page_count': doc.page_count,
            'title': doc.metadata.get('title', ''),
            'author': doc.metadata.get('author', ''),
            'subject': doc.metadata.get('subject', ''),
            'creator': doc.metadata.get('creator', ''),
            'creation_date': doc.metadata.get('creationDate', ''),
            'modification_date': doc.metadata.get('modDate', '')
        }
        
        for page_num in range(doc.page_count):
            page = doc[page_num]
            page_text = page.get_text()
            text += f"\n--- Page {page_num + 1} ---\n{page_text}"
            
            # Extract page structure
            page_info = {
                'page_number': page_num + 1,
                'text_length': len(page_text),
                'has_images': len(page.get_images()) > 0,
                'image_count': len(page.get_images())
            }
            structure['pages'].append(page_info)
            
            # Extract images info
            for img_index, img in enumerate(page.get_images()):
                structure['images'].append({
                    'page': page_num + 1,
                    'index': img_index,
                    'xref': img[0]
                })
        
        doc.close()
        
        # Also try pdfplumber for table extraction
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    tables = page.extract_tables()
                    if tables:
                        for table_index, table in enumerate(tables):
                            structure['tables'].append({
                                'page': page_num + 1,
                                'table_index': table_index,
                                'rows': len(table),
                                'columns': len(table[0]) if table else 0
                            })
        except Exception as e:
            logger.warning(f"Table extraction failed: {str(e)}")
        
        return {
            'text': text,
            'metadata': metadata,
            'structure': structure,
            'success': True,
            'error': None
        }
    
    def _extract_docx_enhanced(self, file_path: str) -> Dict[str, Any]:
        """Enhanced DOCX extraction with structure analysis"""
        doc = DocxDocument(file_path)
        text = ""
        structure = {'paragraphs': [], 'tables': [], 'images': []}
        
        # Extract paragraphs with style info
        for para_index, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text
            text += para_text + "\n"
            
            structure['paragraphs'].append({
                'index': para_index,
                'text_length': len(para_text),
                'style': paragraph.style.name if paragraph.style else 'Normal'
            })
        
        # Extract tables
        for table_index, table in enumerate(doc.tables):
            structure['tables'].append({
                'index': table_index,
                'rows': len(table.rows),
                'columns': len(table.columns)
            })
            
            # Add table text to main text
            text += f"\n--- Table {table_index + 1} ---\n"
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                text += row_text + "\n"
        
        # Basic metadata
        metadata = {
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables),
            'title': doc.core_properties.title or '',
            'author': doc.core_properties.author or '',
            'subject': doc.core_properties.subject or '',
            'created': str(doc.core_properties.created) if doc.core_properties.created else '',
            'modified': str(doc.core_properties.modified) if doc.core_properties.modified else ''
        }
        
        return {
            'text': text,
            'metadata': metadata,
            'structure': structure,
            'success': True,
            'error': None
        }
    
    def _extract_image_enhanced(self, file_path: str) -> Dict[str, Any]:
        """Enhanced image OCR with metadata and multilingual support"""
        try:
            image = Image.open(file_path)
            
            # Try to detect language first
            detected_lang = pytesseract.image_to_osd(image, output_type=pytesseract.Output.DICT)
            lang_script = detected_lang.get('script', 'Latin')
            
            # Set language based on script detection
            # Tesseract language codes: eng (English), hin (Hindi), san (Sanskrit), etc.
            lang_codes = {
                'Latin': 'eng',
                'Arabic': 'ara',
                'Chinese': 'chi_sim',
                'Japanese': 'jpn',
                'Korean': 'kor',
                'Devanagari': 'hin',
                'Armenian': 'arm',
                'Bengali': 'ben',
                'Cyrillic': 'rus',
                'Ethiopic': 'amh',
                'Greek': 'ell',
                'Gujarati': 'guj',
                'Gurmukhi': 'pan',
                'Kannada': 'kan',
                'Malayalam': 'mal',
                'Myanmar': 'mya',
                'Oriya': 'ori',
                'Sinhala': 'sin',
                'Tamil': 'tam',
                'Telugu': 'tel',
                'Thai': 'tha'
            }
            
            # Default to English if script not in mapping
            detected_language = lang_codes.get(lang_script, 'eng')
            
            # Perform OCR with detected language
            text = pytesseract.image_to_string(image, lang=detected_language)
            
            metadata = {
                'format': image.format,
                'mode': image.mode,
                'size': image.size,
                'width': image.width,
                'height': image.height,
                'detected_script': lang_script,
                'ocr_language': detected_language
            }
            
            # Get OCR confidence data
            ocr_data = pytesseract.image_to_data(image, lang=detected_language, output_type=pytesseract.Output.DICT)
            confidence_scores = [int(conf) for conf in ocr_data['conf'] if int(conf) > 0]
            avg_confidence = sum(confidence_scores) / len(confidence_scores) if confidence_scores else 0
            
            structure = {
                'ocr_confidence': avg_confidence,
                'word_count': len([word for word in ocr_data['text'] if word.strip()]),
                'detected_languages': detected_language
            }
            
            return {
                'text': text,
                'metadata': metadata,
                'structure': structure,
                'success': True,
                'error': None
            }
            
        except Exception as e:
            # Fallback to English if language detection fails
            try:
                image = Image.open(file_path)
                text = pytesseract.image_to_string(image, lang='eng')
                
                metadata = {
                    'format': image.format,
                    'mode': image.mode,
                    'size': image.size,
                    'width': image.width,
                    'height': image.height,
                    'detected_script': 'Unknown',
                    'ocr_language': 'eng',
                    'fallback_used': True
                }
                
                ocr_data = pytesseract.image_to_data(image, lang='eng', output_type=pytesseract.Output.DICT)
                confidence_scores = [int(conf) for conf in ocr_data['conf'] if int(conf) > 0]
                avg_confidence = sum(confidence_scores) / len(confidence_scores) if confidence_scores else 0
                
                structure = {
                    'ocr_confidence': avg_confidence,
                    'word_count': len([word for word in ocr_data['text'] if word.strip()]),
                    'detected_languages': 'eng'
                }
                
                return {
                    'text': text,
                    'metadata': metadata,
                    'structure': structure,
                    'success': True,
                    'error': None
                }
            except Exception as fallback_e:
                return {
                    'text': '',
                    'metadata': {},
                    'structure': {},
                    'success': False,
                    'error': f"Primary error: {str(e)}, Fallback error: {str(fallback_e)}"
                }
    
    def ai_question_answer(self, document_text: str, question: str, context: Dict = None) -> Dict[str, Any]:
        """Advanced AI-powered Q&A with document analysis and reasoning
        
        Args:
            document_text: The full text content of the document
            question: The question to answer about the document
            context: Additional context (e.g., document metadata, user info)
            
        Returns:
            Dict containing the answer and metadata
        """
        try:
            # Enhanced system prompt with reasoning capabilities
            system_prompt = """You are an expert AI assistant specialized in analyzing and answering questions 
            about document content with maximum accuracy and detail.
            
INSTRUCTIONS:
1. Carefully analyze the entire document content provided
2. For factual questions, provide direct quotes from the text when possible
3. For analytical questions, provide well-reasoned answers based on the content
4. If the answer isn't in the document, clearly state this
5. For complex questions, break down your answer into logical sections
6. Include relevant statistics, dates, or figures when available
7. Maintain a professional and objective tone
8. If multiple interpretations are possible, present them with their likelihood

OUTPUT FORMAT (JSON):
{
    "answer": "Your detailed, well-structured answer",
    "confidence": 0.0-1.0,
    "sources": ["Relevant document sections or quotes"],
    "key_points": ["Main supporting points from the document"],
    "needs_clarification": true/false,
    "suggested_follow_ups": ["Related questions that might be helpful"]
}"""
            
            # Build the prompt with document and question
            prompt = f"""DOCUMENT CONTENT:
{document_text}

QUESTION: {question}

Please analyze the document and provide a comprehensive answer to the question. 
If the information is not present in the document, clearly state this.

Your response should be in the specified JSON format."""
            
            # Add context if available
            if context:
                prompt += f"\n\nADDITIONAL CONTEXT:\n{json.dumps(context, indent=2)}"
            
            # Call the AI model with enhanced parameters
            response = requests.post(
                f"{self.gemini_api_url}?key={self.gemini_api_key}",
                json={
                    "contents": [{
                        "parts": [
                            {"text": system_prompt},
                            {"text": prompt}
                        ]
                    }],
                    "generationConfig": {
                        "temperature": 0.2,  # Lower temperature for more focused answers
                        "topP": 0.9,
                        "topK": 40,
                        "maxOutputTokens": 4096,  # Increased for more detailed answers
                    },
                    "safetySettings": [
                        {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
                        {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
                        {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"}
                    ]
                },
                timeout=45  # Increased timeout for complex documents
            )
            
            if response.status_code == 200:
                result = response.json()
                if 'candidates' in result and result['candidates']:
                    try:
                        # Extract and clean the response
                        response_text = result['candidates'][0]['content']['parts'][0]['text']
                        response_text = re.sub(r'```json\n|```', '', response_text).strip()
                        
                        # Parse the JSON response
                        answer_data = json.loads(response_text)
                        
                        # Validate and enhance the response
                        if 'answer' not in answer_data:
                            raise ValueError("Invalid response format: missing 'answer' field")
                            
                        # Ensure required fields exist with defaults
                        answer_data.setdefault('confidence', 0.8)
                        answer_data.setdefault('sources', ['Full document'])
                        answer_data.setdefault('key_points', [])
                        answer_data.setdefault('needs_clarification', False)
                        answer_data.setdefault('suggested_follow_ups', [])
                        
                        # Add success flag
                        answer_data['success'] = True
                        
                        return answer_data
                        
                    except (json.JSONDecodeError, KeyError, ValueError) as e:
                        logger.error(f"Failed to parse AI response: {str(e)}")
                        # Fallback to a structured error response
                        return {
                            'success': True,
                            'answer': f"I found this information in the document:\n\n{response_text}",
                            'confidence': 0.6,
                            'sources': ['Full document'],
                            'key_points': [],
                            'needs_clarification': True,
                            'suggested_follow_ups': [
                                "Could you clarify your question?",
                                "Would you like me to look for specific information?"
                            ]
                        }
            
            # Handle API errors
            error_msg = f"API Error: {response.status_code}"
            if response.text:
                try:
                    error_data = response.json()
                    error_msg = error_data.get('error', {}).get('message', str(error_data))
                except:
                    error_msg = response.text[:500]  # Truncate long error messages
            
            logger.error(f"AI Q&A API error: {error_msg}")
            return {
                'success': False,
                'error': f"Failed to get AI response: {error_msg}",
                'question': question,
                'suggested_actions': [
                    "Try rephrasing your question",
                    "Check if the document contains the relevant information",
                    "Try a more specific or different question"
                ]
            }
                
        except Exception as e:
            logger.error(f"Error in AI Q&A: {str(e)}", exc_info=True)
            return {
                'success': False,
                'error': f"An unexpected error occurred: {str(e)}",
                'question': question,
                'suggested_actions': [
                    "Please try again later",
                    "Contact support if the issue persists"
                ]
            }
    
    def smart_edit_content(self, document_text: str, edit_instruction: str, 
                          document_metadata: Dict = None) -> Dict[str, Any]:
        """AI-powered smart editing with document structure awareness and change tracking
        
        Args:
            document_text: The original text content to be edited
            edit_instruction: User's instruction for the edit
            document_metadata: Optional metadata about the document (type, structure, etc.)
            
        Returns:
            Dict containing the edit results or error information
            {
                'success': bool,
                'edited_content': str,
                'changes_made': List[Dict],
                'confidence': float,
                'notes': str,
                'needs_review': bool,
                'suggested_next_steps': List[str],
                'error': Optional[str]
            }
        """
        try:
            # Enhanced system prompt with document awareness
            system_prompt = """You are an expert content editor with deep expertise in document processing.
            
DOCUMENT EDITING INSTRUCTIONS:
1. Carefully analyze the original content and edit instruction
2. Apply ONLY the specific changes requested in the instruction
3. Preserve the core meaning, factual accuracy, and document structure
4. Maintain consistent formatting and style throughout the document
5. For legal/formal documents, preserve precise language and formatting
6. For creative content, enhance style while maintaining the original voice
7. Clearly document all changes made
8. If the instruction is unclear or cannot be safely applied, explain why

OUTPUT FORMAT (JSON):
{
    "edited_content": "The complete edited document content",
    "changes_made": [
        {
            "type": "addition|deletion|modification|formatting",
            "location": "Section/paragraph reference",
            "description": "Description of the change",
            "before": "Original text (for modifications/deletions)",
            "after": "New text (for modifications/additions)"
        }
    ],
    "confidence": 0.0-1.0,
    "notes": "Any additional context or explanations",
    "needs_review": true/false,
    "suggested_next_steps": ["Suggestions for further improvements"]
}"""

            # Build the prompt with document context and edit instruction
            prompt_parts = [
                "DOCUMENT TO EDIT:",
                document_text[:12000],  # Limit size to avoid token limits
                "\nEDIT INSTRUCTION:",
                edit_instruction
            ]
            
            # Add document metadata if available
            if document_metadata:
                prompt_parts.extend([
                    "\nDOCUMENT METADATA:",
                    json.dumps(document_metadata, indent=2)
                ])
            
            prompt = "\n".join(prompt_parts)
            
            # Call the AI model with enhanced parameters
            response = requests.post(
                f"{self.gemini_api_url}?key={self.gemini_api_key}",
                json={
                    "contents": [{
                        "parts": [
                            {"text": system_prompt},
                            {"text": prompt}
                        ]
                    }],
                    "generationConfig": {
                        "temperature": 0.3,  # Lower temperature for more precise edits
                        "topP": 0.9,
                        "topK": 40,
                        "maxOutputTokens": 4096,
                    },
                    "safetySettings": [
                        {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
                        {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
                        {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"}
                    ]
                },
                timeout=60  # Increased timeout for complex edits
            )
            
            if response.status_code == 200:
                result = response.json()
                if 'candidates' in result and result['candidates']:
                    try:
                        # Extract and clean the response
                        response_text = result['candidates'][0]['content']['parts'][0]['text']
                        response_text = re.sub(r'```json\n|```', '', response_text).strip()
                        
                        # Parse the JSON response
                        edit_result = json.loads(response_text)
                        
                        # Validate the response
                        if 'edited_content' not in edit_result:
                            raise ValueError("Invalid response: missing 'edited_content' field")
                            
                        # Ensure required fields with defaults
                        edit_result.setdefault('changes_made', [{"type": "modification", "description": "Applied edit based on instruction"}])
                        edit_result.setdefault('confidence', 0.8)
                        edit_result.setdefault('notes', '')
                        edit_result.setdefault('needs_review', False)
                        edit_result.setdefault('suggested_next_steps', [])
                        
                        # Add original content and instruction to the result
                        edit_result.update({
                            'success': True,
                            'original_content': document_text,
                            'edit_instruction': edit_instruction,
                            'timestamp': datetime.utcnow().isoformat()
                        })
                        
                        return edit_result
                        
                    except (json.JSONDecodeError, KeyError, ValueError) as e:
                        logger.error(f"Failed to parse AI edit response: {str(e)}")
                        # Fallback to a structured response with the raw edit
                        return {
                            'success': True,
                            'original_content': document_text,
                            'edited_content': response_text,
                            'edit_instruction': edit_instruction,
                            'changes_made': [{
                                'type': 'modification',
                                'description': 'Applied edit based on instruction (fallback)',
                                'notes': 'The AI response format was unexpected. Review carefully.'
                            }],
                            'confidence': 0.6,
                            'needs_review': True,
                            'suggested_next_steps': [
                                "Review the changes for accuracy",
                                "Consider rephrasing the edit instruction if needed"
                            ]
                        }
            
            # Handle API errors
            error_msg = f"API Error: {response.status_code}"
            if response.text:
                try:
                    error_data = response.json()
                    error_msg = error_data.get('error', {}).get('message', str(error_data))
                except:
                    error_msg = response.text[:500]
            
            logger.error(f"Smart Edit API error: {error_msg}")
            return {
                'success': False,
                'error': f"Failed to process edit: {error_msg}",
                'edit_instruction': edit_instruction,
                'needs_review': True,
                'suggested_next_steps': [
                    "Check your internet connection",
                    "Verify the API key is valid",
                    "Try again with a simpler edit instruction"
                ]
            }
                
        except Exception as e:
            logger.error(f"Error in smart_edit_content: {str(e)}", exc_info=True)
            return {
                'success': False,
                'error': f"An unexpected error occurred: {str(e)}",
                'edit_instruction': edit_instruction,
                'needs_review': True,
                'suggested_next_steps': [
                    "Try again with a different edit instruction",
                    "Check the server logs for more details"
                ]
            }
    
    def convert_document_format(self, content: str, source_format: str, target_format: str, 
                              metadata: Dict = None) -> Dict[str, Any]:
        """Convert document to different formats"""
        try:
            if target_format.lower() == 'json':
                return self._convert_to_json(content, metadata)
            elif target_format.lower() == 'pdf':
                return self._convert_to_pdf(content)
            elif target_format.lower() == 'docx':
                return self._convert_to_docx(content)
            elif target_format.lower() == 'markdown':
                return self._convert_to_markdown(content)
            elif target_format.lower() == 'html':
                return self._convert_to_html(content)
            elif target_format.lower() == 'txt':
                return self._convert_to_txt(content)
            else:
                return {
                    'success': False,
                    'error': f"Unsupported target format: {target_format}"
                }
                
        except Exception as e:
            logger.error(f"Error converting to {target_format}: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def _convert_to_json(self, content: str, metadata: Dict = None) -> Dict[str, Any]:
        """Convert content to structured JSON with enhanced organization"""
        try:
            # Parse content into structured format with better organization
            lines = content.split('\n')
            paragraphs = [line.strip() for line in lines if line.strip()]
            
            # Split content into sections if possible
            sections = self._identify_sections(content)
            
            # Extract key information
            key_info = self._extract_key_information(content)
            
            structured_data = {
                'document_info': {
                    'total_characters': len(content),
                    'total_words': len(content.split()),
                    'total_paragraphs': len(paragraphs),
                    'total_sections': len(sections),
                    'extraction_timestamp': datetime.utcnow().isoformat()
                },
                'metadata': metadata or {},
                'content': {
                    'full_text': content,
                    'paragraphs': paragraphs,
                    'sections': sections,
                    'key_information': key_info
                }
            }
            
            # Create temporary file
            temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False)
            json.dump(structured_data, temp_file, indent=2, ensure_ascii=False)
            temp_file.close()
            
            return {
                'success': True,
                'file_path': temp_file.name,
                'format': 'json',
                'data': structured_data
            }
        except Exception as e:
            logger.error(f"Error in JSON conversion: {str(e)}")
            # Fallback to simple conversion
            lines = content.split('\n')
            paragraphs = [line.strip() for line in lines if line.strip()]
            
            structured_data = {
                'document_info': {
                    'total_characters': len(content),
                    'total_words': len(content.split()),
                    'total_paragraphs': len(paragraphs),
                    'extraction_timestamp': datetime.utcnow().isoformat()
                },
                'metadata': metadata or {},
                'content': {
                    'full_text': content,
                    'paragraphs': paragraphs,
                    'summary': content[:500] + '...' if len(content) > 500 else content
                }
            }
            
            # Create temporary file
            temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False)
            json.dump(structured_data, temp_file, indent=2, ensure_ascii=False)
            temp_file.close()
            
            return {
                'success': True,
                'file_path': temp_file.name,
                'format': 'json',
                'data': structured_data
            }
    
    def _convert_to_pdf(self, content: str) -> Dict[str, Any]:
        """Convert content to PDF"""
        temp_file = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
        temp_file.close()
        
        # Create PDF using ReportLab
        doc = SimpleDocTemplate(temp_file.name, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                p = Paragraph(para.strip(), styles['Normal'])
                story.append(p)
                story.append(Spacer(1, 12))
        
        doc.build(story)
        
        return {
            'success': True,
            'file_path': temp_file.name,
            'format': 'pdf'
        }
    
    def _convert_to_docx(self, content: str) -> Dict[str, Any]:
        """Convert content to DOCX"""
        doc = DocxDocument()
        
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        doc.save(temp_file.name)
        temp_file.close()
        
        return {
            'success': True,
            'file_path': temp_file.name,
            'format': 'docx'
        }
    
    def _convert_to_markdown(self, content: str) -> Dict[str, Any]:
        """Convert content to Markdown"""
        # Basic conversion - could be enhanced with structure detection
        markdown_content = content
        
        temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8')
        temp_file.write(markdown_content)
        temp_file.close()
        
        return {
            'success': True,
            'file_path': temp_file.name,
            'format': 'markdown'
        }
    
    def _convert_to_html(self, content: str) -> Dict[str, Any]:
        """Convert content to HTML with enhanced structure and styling"""
        try:
            # Identify sections in the content
            sections = self._identify_sections(content)
            
            html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>AI DocTransform - Converted Document</title>
    <style>
        body {{ 
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; 
            line-height: 1.6; 
            margin: 40px auto; 
            max-width: 800px; 
            padding: 20px;
            background-color: #f9f9f9;
            color: #333;
        }}
        header {{ 
            text-align: center; 
            padding: 20px 0; 
            border-bottom: 2px solid #007bff; 
            margin-bottom: 30px;
        }}
        h1 {{ color: #007bff; }}
        h2 {{ color: #0056b3; margin-top: 30px; }}
        h3 {{ color: #004085; }}
        p {{ margin-bottom: 16px; text-align: justify; }}
        .section {{ 
            background: white; 
            padding: 20px; 
            margin: 20px 0; 
            border-radius: 5px; 
            box-shadow: 0 2px 4px rgba(0,0,0,0.1); 
        }}
        .metadata {{ 
            background: #e9ecef; 
            padding: 15px; 
            border-left: 4px solid #007bff; 
            margin: 20px 0; 
            font-size: 0.9em; 
        }}
        footer {{ 
            text-align: center; 
            margin-top: 40px; 
            padding-top: 20px; 
            border-top: 1px solid #ddd; 
            color: #6c757d; 
            font-size: 0.9em; 
        }}
    </style>
</head>
<body>
    <header>
        <h1>AI DocTransform - Converted Document</h1>
        <p>Processed on {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC</p>
    </header>
"""
            
            # Add sections if identified
            if sections and len(sections) > 1:
                for i, section in enumerate(sections):
                    title = section.get('title', f'Section {i+1}')
                    content = section.get('content', '')
                    html_content += f"    <div class=\"section\">\n"
                    html_content += f"        <h2>{title}</h2>\n"
                    # Convert paragraphs within section
                    paragraphs = content.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            html_content += f"        <p>{para.strip()}</p>\n"
                    html_content += "    </div>\n"
            else:
                # Convert paragraphs to HTML
                paragraphs = content.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        html_content += f"    <p>{para.strip()}</p>\n"
            
            html_content += """    <footer>
        <p>Document processed by AI DocTransform - Smart Document Converter & Query Assistant</p>
    </footer>
</body>
</html>"""
            
            temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8')
            temp_file.write(html_content)
            temp_file.close()
            
            return {
                'success': True,
                'file_path': temp_file.name,
                'format': 'html'
            }
        except Exception as e:
            logger.error(f"Error in HTML conversion: {str(e)}")
            # Fallback to simple conversion
            html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Converted Document</title>
    <style>
        body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; }}
        p {{ margin-bottom: 16px; }}
    </style>
</head>
<body>
"""
            
            # Convert paragraphs to HTML
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    html_content += f"    <p>{para.strip()}</p>\n"
            
            html_content += """</body>
</html>"""
            
            temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8')
            temp_file.write(html_content)
            temp_file.close()
            
            return {
                'success': True,
                'file_path': temp_file.name,
                'format': 'html'
            }
    
    def _identify_sections(self, content: str) -> list:
        """Identify sections in document content based on headers and structure"""
        try:
            sections = []
            lines = content.split('\n')
            
            # Common header patterns
            header_patterns = [
                r'^(#+)\s+(.+)',  # Markdown headers
                r'^([A-Z][A-Z0-9\s]{2,}:)$',  # ALL CAPS headers
                r'^([A-Z][a-z]+\s+[A-Z][a-z]+):$',  # Title Case headers
                r'^(\d+\.\s+.+)$',  # Numbered sections
                r'^([IVX]+\.\s+.+)$'  # Roman numeral sections
            ]
            
            current_section = None
            current_content = []
            
            for line in lines:
                is_header = False
                
                # Check for headers
                for pattern in header_patterns:
                    if re.match(pattern, line.strip()):
                        is_header = True
                        break
                
                # If we found a header
                if is_header:
                    # Save previous section if it exists
                    if current_section:
                        sections.append({
                            'title': current_section,
                            'content': '\n'.join(current_content).strip()
                        })
                    
                    # Start new section
                    current_section = line.strip()
                    current_content = []
                else:
                    # Add line to current section
                    if current_section or len(current_content) > 0:
                        current_content.append(line)
            
            # Add final section
            if current_section:
                sections.append({
                    'title': current_section,
                    'content': '\n'.join(current_content).strip()
                })
            
            return sections
        except Exception as e:
            logger.error(f"Error identifying sections: {str(e)}")
            return []
    
    def _extract_key_information(self, content: str) -> dict:
        """Extract key information like dates, numbers, and important entities"""
        try:
            key_info = {
                'dates': [],
                'numbers': [],
                'email_addresses': [],
                'phone_numbers': [],
                'important_entities': []
            }
            
            # Extract dates (various formats)
            date_patterns = [
                r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',  # MM/DD/YYYY or MM-DD-YYYY
                r'\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b',     # YYYY/MM/DD or YYYY-MM-DD
                r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{2,4}\b'  # Month DD, YYYY
            ]
            
            for pattern in date_patterns:
                dates = re.findall(pattern, content, re.IGNORECASE)
                key_info['dates'].extend(dates)
            
            # Extract numbers (excluding dates)
            number_pattern = r'\b\d+(?:[,.]\d+)*\b'
            numbers = re.findall(number_pattern, content)
            # Filter out likely dates
            for num in numbers:
                if not re.match(r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}', num) and \
                   not re.match(r'\d{4}[/-]\d{1,2}[/-]\d{1,2}', num):
                    key_info['numbers'].append(num)
            
            # Extract email addresses
            email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
            key_info['email_addresses'] = re.findall(email_pattern, content)
            
            # Extract phone numbers
            phone_patterns = [
                r'\b\d{3}-\d{3}-\d{4}\b',
                r'\b\(\d{3}\)\s*\d{3}-\d{4}\b',
                r'\b\d{3}\.\d{3}\.\d{4}\b'
            ]
            
            for pattern in phone_patterns:
                phones = re.findall(pattern, content)
                key_info['phone_numbers'].extend(phones)
            
            # Extract potential important entities (capitalized phrases)
            entity_pattern = r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3}\b'
            entities = re.findall(entity_pattern, content)
            # Filter common words
            common_words = {'The', 'And', 'For', 'With', 'From', 'This', 'That', 'Have', 'Were', 'Where', 'When', 'What', 'Who', 'Why', 'How'}
            key_info['important_entities'] = [e for e in entities if e.split()[0] not in common_words]
            
            return key_info
        except Exception as e:
            logger.error(f"Error extracting key information: {str(e)}")
            return {}
    
    def _convert_to_txt(self, content: str) -> Dict[str, Any]:
        """Convert content to plain text"""
        temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8')
        temp_file.write(content)
        temp_file.close()
        
        return {
            'success': True,
            'file_path': temp_file.name,
            'format': 'txt'
        }
    
    def generate_summary(self, content: str, summary_type: str = 'brief') -> Dict[str, Any]:
        """Generate AI-powered summary of document content with enhanced accuracy"""
        try:
            # Enhanced instructions based on summary type
            if summary_type == 'brief':
                instruction = "Provide a concise summary (2-3 sentences) capturing ONLY the main points."
                format_instruction = "Format as continuous text."
            elif summary_type == 'detailed':
                instruction = "Provide a comprehensive summary covering all major points, key details, and important facts."
                format_instruction = "Format as continuous text with clear paragraph structure."
            elif summary_type == 'bullet':
                instruction = "Provide a structured bullet-point summary of the key points and main ideas."
                format_instruction = "Format as a clear bullet list with 5-10 key points."
            elif summary_type == 'executive':
                instruction = "Provide an executive summary highlighting critical insights, conclusions, and recommendations."
                format_instruction = "Format as 3-5 concise paragraphs focusing on key takeaways."
            else:
                instruction = "Provide a comprehensive summary covering all major points, key details, and important facts."
                format_instruction = "Format as continuous text with clear paragraph structure."
            
            # Enhanced system prompt for better summarization accuracy
            system_prompt = f"""You are an expert document summarization specialist with exceptional analytical skills.
            
INSTRUCTIONS:
1. {instruction}
2. Extract ONLY information present in the original document
3. Maintain factual accuracy and avoid adding assumptions
4. {format_instruction}
5. Preserve key numbers, dates, and specific facts
6. Eliminate redundant information while maintaining completeness"""
            
            user_prompt = f"""DOCUMENT CONTENT TO SUMMARIZE:
{content[:7000]}  # Limit for API

SUMMARY TYPE: {summary_type}

Please provide the requested summary following the format instructions above."""
            
            headers = {
                'Content-Type': 'application/json',
            }
            
            data = {
                "contents": [{
                    "parts": [{
                        "text": f"{system_prompt}\n\n{user_prompt}"
                    }]
                }],
                "generationConfig": {
                    "temperature": 0.3,  # Lower temperature for more factual summaries
                    "maxOutputTokens": 2048,
                    "topK": 40,
                    "topP": 0.95
                },
                "safetySettings": [
                    {
                        "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
                        "threshold": "BLOCK_ONLY_HIGH"
                    }
                ]
            }
            
            response = requests.post(
                f"{self.gemini_api_url}?key={self.gemini_api_key}",
                headers=headers,
                json=data,
                timeout=45  # Increased timeout for better processing
            )
            
            if response.status_code == 200:
                result = response.json()
                
                # Check if response has content
                if 'candidates' in result and len(result['candidates']) > 0:
                    summary = result['candidates'][0]['content']['parts'][0]['text']
                    
                    # Calculate quality metrics
                    original_length = len(content)
                    summary_length = len(summary)
                    compression_ratio = summary_length / original_length if original_length > 0 else 0
                    
                    # Determine quality based on compression ratio
                    if summary_type == 'brief':
                        quality = 'high' if 0.05 <= compression_ratio <= 0.2 else 'medium'
                    elif summary_type == 'detailed':
                        quality = 'high' if 0.15 <= compression_ratio <= 0.4 else 'medium'
                    else:
                        quality = 'high' if 0.1 <= compression_ratio <= 0.3 else 'medium'
                    
                    return {
                        'success': True,
                        'summary': summary,
                        'summary_type': summary_type,
                        'original_length': original_length,
                        'summary_length': summary_length,
                        'compression_ratio': compression_ratio,
                        'quality': quality,
                        'timestamp': datetime.utcnow().isoformat()
                    }
                else:
                    return {
                        'success': False,
                        'error': 'No valid response generated from AI model',
                        'summary_type': summary_type
                    }
            else:
                return {
                    'success': False,
                    'error': f"API Error: {response.status_code} - {response.text}",
                    'summary_type': summary_type
                }
                
        except Exception as e:
            logger.error(f"Error generating summary: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'summary_type': summary_type
            }
    
    def answer_question(self, document_text: str, question: str, document_id: int = None, 
                       user_id: int = None, chat_history: list = None) -> tuple[str, str]:
        """Answer a question about a document with chat history context
        
        Args:
            document_text: The text content of the document
            question: The user's question about the document
            document_id: Optional ID of the document in database
            user_id: Optional ID of the user asking the question
            chat_history: List of previous chat messages for context
            
        Returns:
            tuple: (answer_text, job_uuid) or (error_message, None)
        """
        try:
            # Build context with document and chat history
            context = f"Document Content:\n{document_text}"
            
            if chat_history:
                # Add last 5 messages for context
                context += "\n\nChat History:"
                for msg in chat_history[-5:]:
                    role = "User" if msg.get('role') == 'user' else "AI"
                    context += f"\n{role}: {msg.get('content', '')}"
            
            # Get AI response with context
            result = self.ai_question_answer(context, question)
            
            if result.get('success'):
                answer = result['answer']
                
                # Save processing job with metadata
                job = ProcessingJob(
                    job_type='qa',
                    input_text=question,
                    output_text=answer,
                    document_id=document_id,
                    user_id=user_id,
                    status='completed',
                    metadata={
                        'model': 'gemini-pro',
                        'context_length': len(context),
                        'chat_history_used': len(chat_history) if chat_history else 0
                    }
                )
                db.session.add(job)
                db.session.commit()
                
                return answer, job.uuid
                
            return result.get('error', 'Failed to generate answer'), None
            
        except Exception as e:
            logger.error(f"Error in answer_question: {str(e)}", exc_info=True)
            return f"An error occurred while processing your question: {str(e)}", None
